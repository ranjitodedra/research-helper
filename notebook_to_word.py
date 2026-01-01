#!/usr/bin/env python3
"""
Convert Jupyter Notebook (.ipynb) to Microsoft Word (.docx) file.
"""

import nbformat
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import re
import os
import tempfile
from pathlib import Path
from PIL import Image
import io


def extract_base64_images(markdown_text):
    """Extract base64 encoded PNG images from markdown."""
    images = []
    # Pattern for base64 images: ![alt](data:image/png;base64,...)
    pattern = r'!\[.*?\]\(data:image/png;base64,([A-Za-z0-9+/=]+)\)'
    matches = re.finditer(pattern, markdown_text)
    
    for match in matches:
        base64_data = match.group(1)
        try:
            image_data = base64.b64decode(base64_data)
            images.append(image_data)
        except Exception:
            continue
    
    return images


def extract_cell_images(cell):
    """Extract PNG images from cell outputs."""
    images = []
    if cell.cell_type == 'code' and 'outputs' in cell:
        for output in cell.get('outputs', []):
            if output.get('output_type') == 'display_data':
                data = output.get('data', {})
                if 'image/png' in data:
                    try:
                        image_data = base64.b64decode(data['image/png'])
                        images.append(image_data)
                    except Exception:
                        continue
    return images


def resize_image_for_word(image_path, max_width=6.5, max_height=9.0):
    """Resize image to fit within Word page dimensions."""
    with Image.open(image_path) as img:
        width, height = img.size
        
        # Calculate scaling factor
        width_inches = width / 96  # Assuming 96 DPI
        height_inches = height / 96
        
        scale_w = max_width / width_inches if width_inches > max_width else 1.0
        scale_h = max_height / height_inches if height_inches > max_height else 1.0
        scale = min(scale_w, scale_h, 1.0)  # Don't upscale
        
        if scale < 1.0:
            new_width = int(width * scale)
            new_height = int(height * scale)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            img.save(image_path)
        
        return image_path


def save_image_temp(image_data, temp_dir, index):
    """Save image data to temporary file."""
    image_path = os.path.join(temp_dir, f'image_{index}.png')
    with open(image_path, 'wb') as f:
        f.write(image_data)
    return image_path


def add_page_break(doc):
    """Add a page break to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def set_heading_style(paragraph, level):
    """Set paragraph style based on heading level."""
    if level == 1:
        paragraph.style = 'Heading 1'
        paragraph.paragraph_format.space_after = Pt(12)
    elif level == 2:
        paragraph.style = 'Heading 2'
        paragraph.paragraph_format.space_after = Pt(10)
    elif level == 3:
        paragraph.style = 'Heading 3'
        paragraph.paragraph_format.space_after = Pt(8)


def parse_markdown_heading(line):
    """Parse markdown heading and return level and text."""
    line = line.strip()
    if line.startswith('### '):
        return 3, line[4:].strip()
    elif line.startswith('## '):
        return 2, line[3:].strip()
    elif line.startswith('# '):
        return 1, line[2:].strip()
    return None, None


def add_images_to_table(doc, image_paths):
    """Add images to Word document in a 2-column table."""
    if not image_paths:
        return
    
    # Create table with 2 columns
    num_images = len(image_paths)
    num_rows = (num_images + 1) // 2
    
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    
    for idx, image_path in enumerate(image_paths):
        row = idx // 2
        col = idx % 2
        cell = table.rows[row].cells[col]
        
        # Resize image
        resized_path = resize_image_for_word(image_path)
        
        # Add image to cell
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        try:
            run.add_picture(resized_path, width=Inches(3.0))
        except Exception:
            pass


def convert_notebook_to_word(notebook_path, output_path):
    """Convert Jupyter Notebook to Word document."""
    # Read notebook
    with open(notebook_path, 'r', encoding='utf-8') as f:
        nb = nbformat.read(f, as_version=4)
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Temporary directory for images
    with tempfile.TemporaryDirectory() as temp_dir:
        current_h1_images = []
        current_h1_text = []
        is_first_h1 = True
        seen_images = set()
        
        for cell in nb.cells:
            # Extract images from cell outputs
            cell_images = extract_cell_images(cell)
            
            if cell.cell_type == 'markdown':
                # Extract base64 images from markdown
                markdown_images = extract_base64_images(cell.source)
                cell_images.extend(markdown_images)
                
                # Parse markdown content
                lines = cell.source.split('\n')
                for line in lines:
                    level, text = parse_markdown_heading(line)
                    
                    if level == 1:
                        # Process previous H1 section
                        if not is_first_h1:
                            # Add images from previous H1
                            if current_h1_images:
                                image_paths = []
                                for img_data in current_h1_images:
                                    img_hash = hash(img_data)
                                    if img_hash not in seen_images:
                                        seen_images.add(img_hash)
                                        idx = len(seen_images)
                                        img_path = save_image_temp(img_data, temp_dir, idx)
                                        image_paths.append(img_path)
                                
                                if image_paths:
                                    add_images_to_table(doc, image_paths)
                            
                            # Add page break after H1 section
                            add_page_break(doc)
                        
                        # Start new H1 section
                        current_h1_images = []
                        current_h1_text = []
                        is_first_h1 = False
                        
                        # Add H1 heading
                        paragraph = doc.add_paragraph()
                        set_heading_style(paragraph, 1)
                        paragraph.add_run(text)
                    
                    elif level == 2:
                        # Add H2 heading
                        paragraph = doc.add_paragraph()
                        set_heading_style(paragraph, 2)
                        paragraph.add_run(text)
                    
                    elif level == 3:
                        # Add H3 heading
                        paragraph = doc.add_paragraph()
                        set_heading_style(paragraph, 3)
                        paragraph.add_run(text)
                    
                    elif level is None and line.strip():
                        # Regular text (not a heading)
                        # Skip image markdown syntax (already extracted)
                        if not re.match(r'!\[.*?\]\(data:image', line) and not re.match(r'!\[.*?\]\(.*?\)', line):
                            paragraph = doc.add_paragraph(line.strip())
            
            elif cell.cell_type == 'code':
                # Handle code cell outputs (text)
                if 'outputs' in cell:
                    for output in cell.get('outputs', []):
                        if output.get('output_type') == 'stream':
                            text = ''.join(output.get('text', []))
                            if text.strip():
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(text.strip())
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                        elif output.get('output_type') == 'execute_result':
                            data = output.get('data', {})
                            if 'text/plain' in data:
                                text = ''.join(data['text/plain'])
                                if text.strip():
                                    paragraph = doc.add_paragraph()
                                    run = paragraph.add_run(text.strip())
                                    run.font.name = 'Courier New'
                                    run.font.size = Pt(9)
            
            # Collect images for current H1
            for img_data in cell_images:
                img_hash = hash(img_data)
                if img_hash not in seen_images:
                    seen_images.add(img_hash)
                    current_h1_images.append(img_data)
        
        # Process last H1 section
        if current_h1_images:
            image_paths = []
            for img_data in current_h1_images:
                img_hash = hash(img_data)
                if img_hash not in seen_images:
                    seen_images.add(img_hash)
                    idx = len(seen_images)
                    img_path = save_image_temp(img_data, temp_dir, idx)
                    image_paths.append(img_path)
            
            if image_paths:
                add_images_to_table(doc, image_paths)
        
        # Add page break after last H1 if there was content
        if not is_first_h1:
            add_page_break(doc)
    
    # Save document
    doc.save(output_path)
    print(f"Converted {notebook_path} to {output_path}")


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python notebook_to_word.py <notebook_path> [output_path]")
        sys.exit(1)
    
    notebook_path = sys.argv[1]
    
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        output_path = str(Path(notebook_path).with_suffix('.docx'))
    
    convert_notebook_to_word(notebook_path, output_path)

