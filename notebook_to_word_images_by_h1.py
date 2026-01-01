"""
Convert a Jupyter Notebook (.ipynb) into a Word document (.docx) where images are
grouped by Markdown H1 headings.

Strict rules implemented:
- Read notebook with nbformat
- Write Word with python-docx
- Markdown H1 (# Title) is the page title; each H1 starts on a new page
- Everything below an H1 belongs to that H1 until the next H1
- H2/H3 do not create page breaks (they are ignored for pagination/grouping)
- Extract images from:
  - code cell outputs (output.data['image/png'])
  - markdown embedded base64 PNG images
- Preserve original order of images
- Do not duplicate images (global de-dupe across entire document)
- For each H1 section:
  - Put ALL images for that section on ONE page (best-effort: we scale to fit)
  - Layout images in a 2-images-per-row grid using a Word table
  - Resize images uniformly per section to fit the page
  - If an H1 has no images, still create the page with “No images”
- Images before any H1 go under “Untitled”

Output:
- Saved next to the input notebook as: <notebook_name>_images_by_H1.docx

No CLI required: edit NOTEBOOK_PATH below if needed.
"""

from __future__ import annotations

import base64
import hashlib
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import nbformat
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches

try:
    from PIL import Image  # type: ignore
except Exception as e:  # pragma: no cover
    raise RuntimeError(
        "Pillow is required for image sizing. Install with: pip install Pillow"
    ) from e


# -----------------------------
# CONFIG (hardcode path is fine)
# -----------------------------

# Default: convert ho_viz.ipynb in this folder. Change this if you want ch_viz.ipynb.
NOTEBOOK_PATH = Path(__file__).with_name("ch_viz.ipynb")

# If you prefer converting ch_viz.ipynb by default, use:
# NOTEBOOK_PATH = Path(__file__).with_name("ch_viz.ipynb")


# -----------------------------
# Helpers: parsing + extraction
# -----------------------------


H1_RE = re.compile(r"^\s*#\s+(?P<title>.+?)\s*$", re.MULTILINE)

# Markdown image: ![alt](data:image/png;base64,....)
MD_BASE64_IMG_RE = re.compile(
    r"!\[[^\]]*]\(\s*data:image/png;base64,(?P<b64>[A-Za-z0-9+/=\s]+)\s*\)",
    re.IGNORECASE,
)

# HTML <img src="data:image/png;base64,....">
HTML_BASE64_IMG_RE = re.compile(
    r"<img[^>]+src=[\"']data:image/png;base64,(?P<b64>[A-Za-z0-9+/=\s]+)[\"'][^>]*>",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class ExtractedImage:
    section_title: str
    order_index: int  # global order across notebook
    digest: str  # sha256 for de-dupe
    png_bytes: bytes


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _iter_markdown_base64_png_images(md: str) -> List[bytes]:
    """Return list of decoded PNG bytes found in markdown, in appearance order."""
    found: List[bytes] = []
    for regex in (MD_BASE64_IMG_RE, HTML_BASE64_IMG_RE):
        for m in regex.finditer(md):
            b64 = m.group("b64")
            # Remove whitespace/newlines inside base64
            b64_clean = re.sub(r"\s+", "", b64)
            try:
                found.append(base64.b64decode(b64_clean, validate=False))
            except Exception:
                # Skip invalid base64 blocks
                continue
    return found


def _iter_output_png_images(cell: dict) -> List[bytes]:
    """Return list of PNG bytes from a code cell outputs, in appearance order."""
    images: List[bytes] = []
    for out in cell.get("outputs", []) or []:
        data = out.get("data") or {}
        if "image/png" not in data:
            continue
        png_data = data["image/png"]
        # nbformat may store as base64 str or list of lines
        if isinstance(png_data, list):
            png_b64 = "".join(png_data)
        else:
            png_b64 = str(png_data)
        try:
            images.append(base64.b64decode(png_b64, validate=False))
        except Exception:
            continue
    return images


def _get_or_create_section_title(current: Optional[str]) -> str:
    return current if current else "Untitled"


def extract_images_grouped_by_h1(nb_path: Path) -> Tuple[List[str], Dict[str, List[ExtractedImage]]]:
    """
    Parse notebook and return:
    - ordered list of section titles (including 'Untitled' if needed)
    - mapping section_title -> list of ExtractedImage (preserving order, deduped globally)
    """
    nb = nbformat.read(nb_path, as_version=4)

    section_titles_in_order: List[str] = []
    images_by_section: Dict[str, List[ExtractedImage]] = {}

    current_h1: Optional[str] = None
    global_order = 0
    seen_digests: set[str] = set()

    def ensure_section(title: str) -> None:
        if title not in images_by_section:
            images_by_section[title] = []
            section_titles_in_order.append(title)

    for cell in nb.get("cells", []):
        cell_type = cell.get("cell_type")

        # Update current_h1 when we see a markdown H1
        if cell_type == "markdown":
            md = cell.get("source") or ""
            # If multiple H1 in a single markdown cell, treat them in order.
            # We assign images appearing in markdown to the most recent H1 in that cell.
            h1_matches = list(H1_RE.finditer(md))
            if h1_matches:
                # Track last H1 in cell before any images we extract from the same cell
                # We will update current_h1 to the last H1 in the cell, since images in
                # markdown are typically below the heading(s).
                current_h1 = h1_matches[-1].group("title").strip()
                ensure_section(current_h1)

            # Extract base64 markdown images
            for png_bytes in _iter_markdown_base64_png_images(md):
                section = _get_or_create_section_title(current_h1)
                ensure_section(section)
                digest = _sha256(png_bytes)
                if digest in seen_digests:
                    continue
                seen_digests.add(digest)
                images_by_section[section].append(
                    ExtractedImage(section_title=section, order_index=global_order, digest=digest, png_bytes=png_bytes)
                )
                global_order += 1

        # Extract output images from code cells
        if cell_type == "code":
            for png_bytes in _iter_output_png_images(cell):
                section = _get_or_create_section_title(current_h1)
                ensure_section(section)
                digest = _sha256(png_bytes)
                if digest in seen_digests:
                    continue
                seen_digests.add(digest)
                images_by_section[section].append(
                    ExtractedImage(section_title=section, order_index=global_order, digest=digest, png_bytes=png_bytes)
                )
                global_order += 1

    # If there were no H1 at all, but images exist, ensure "Untitled" exists
    if "Untitled" in images_by_section and "Untitled" not in section_titles_in_order:
        section_titles_in_order.insert(0, "Untitled")

    return section_titles_in_order, images_by_section


# -----------------------------
# Word generation
# -----------------------------


def _page_inner_size_inches(doc: Document) -> Tuple[float, float]:
    """Return (inner_width_in, inner_height_in) for the first section."""
    section = doc.sections[0]
    # EMU per inch = 914400; python-docx stores lengths as EMU ints
    emu_per_in = 914400
    inner_w = (section.page_width - section.left_margin - section.right_margin) / emu_per_in
    inner_h = (section.page_height - section.top_margin - section.bottom_margin) / emu_per_in
    return float(inner_w), float(inner_h)


def _compute_target_box_per_image(
    doc: Document,
    num_images: int,
    title_lines: int = 1,
) -> Tuple[float, float]:
    """
    Compute a per-image (max_width_in, max_height_in) box so that a 2-column grid
    can fit on one page (best effort).
    """
    inner_w, inner_h = _page_inner_size_inches(doc)

    cols = 2
    rows = (num_images + cols - 1) // cols

    # Rough space budgeting:
    # - title: ~0.5 inch per line (approx)
    # - some spacing between title and table + table cell padding
    title_h = 0.5 * max(1, title_lines)
    overhead = title_h + 0.4  # extra whitespace budget

    available_h = max(1.0, inner_h - overhead)
    available_w = inner_w

    # Add a small gutter between the 2 columns
    gutter = 0.2
    cell_w = (available_w - gutter) / cols
    cell_h = available_h / max(1, rows)

    # Leave a little padding inside the cell
    pad = 0.05
    return max(0.5, cell_w - pad), max(0.5, cell_h - pad)


def _save_png(tmp_dir: Path, digest: str, png_bytes: bytes) -> Path:
    path = tmp_dir / f"{digest}.png"
    path.write_bytes(png_bytes)
    return path


def _scaled_dims_in_inches(img_path: Path, max_w_in: float, max_h_in: float) -> Tuple[float, float]:
    """Scale image to fit within (max_w_in, max_h_in) keeping aspect ratio."""
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    if w_px <= 0 or h_px <= 0:
        return max_w_in, max_h_in

    # Convert to relative scale (we only need ratios)
    scale = min(max_w_in / w_px, max_h_in / h_px)
    # Avoid zero
    scale = max(scale, 0.0001)
    return w_px * scale, h_px * scale


def write_docx_images_by_h1(
    nb_path: Path,
    out_path: Path,
    section_titles: List[str],
    images_by_section: Dict[str, List[ExtractedImage]],
) -> None:
    doc = Document()

    for idx, title in enumerate(section_titles):
        # Each H1 must start on a new page (except the first page, which is already new).
        if idx != 0:
            doc.add_page_break()

        # Title (H1)
        p = doc.add_paragraph(title)
        try:
            p.style = "Heading 1"
        except Exception:
            # Fallback if style missing
            p.runs[0].bold = True

        imgs = images_by_section.get(title, [])
        if not imgs:
            doc.add_paragraph("No images")
            continue

        # Compute a per-image bounding box to fit all images in one page.
        max_w_in, max_h_in = _compute_target_box_per_image(doc, num_images=len(imgs), title_lines=1)

        cols = 2
        rows = (len(imgs) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = False

        # Fill the table left-to-right, top-to-bottom
        with tempfile.TemporaryDirectory() as td:
            tmp_dir = Path(td)

            for i, img in enumerate(imgs):
                r = i // cols
                c = i % cols
                cell = table.cell(r, c)
                cell_par = cell.paragraphs[0]

                img_path = _save_png(tmp_dir, img.digest, img.png_bytes)
                w_in, h_in = _scaled_dims_in_inches(img_path, max_w_in, max_h_in)

                run = cell_par.add_run()
                run.add_picture(img_path.as_posix(), width=Inches(w_in), height=Inches(h_in))

            # If odd number of images, clear the final empty cell’s paragraph text
            if len(imgs) % cols == 1:
                last_cell = table.cell(rows - 1, 1)
                for par in last_cell.paragraphs:
                    par.text = ""

        # After finishing one H1 section, insert a page break (handled at top of next section).
        # We avoid adding an extra trailing blank page at end.

    doc.save(out_path)


def main() -> None:
    nb_path = NOTEBOOK_PATH
    if not nb_path.exists():
        raise FileNotFoundError(f"Notebook not found: {nb_path.resolve()}")

    section_titles, images_by_section = extract_images_grouped_by_h1(nb_path)

    out_path = nb_path.with_name(f"{nb_path.stem}_images_by_H1.docx")
    write_docx_images_by_h1(nb_path, out_path, section_titles, images_by_section)

    print(f"Saved: {out_path.resolve()}")


if __name__ == "__main__":
    main()


